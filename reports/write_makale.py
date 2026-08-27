#!/usr/bin/env python3
"""Rewrite makale.docx for the locked Ridge+B+AR(1) delivery and the five-model talk order."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "makale.docx"
NAVY = RGBColor(0x1F, 0x2A, 0x44)
GRAY = RGBColor(0x44, 0x44, 0x44)


def set_run_font(run, *, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def fmt(p, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=8, first=0):
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    if first:
        pf.first_line_indent = Cm(first)


def shade_cell(cell, fill="F2F4F8"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


class Doc:
    def __init__(self):
        self.doc = Document()
        sec = self.doc.sections[0]
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.2)

    def title(self, text: str) -> None:
        p = self.doc.add_paragraph()
        fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
        r = p.add_run(text)
        set_run_font(r, size=18, bold=True, color=NAVY)

    def h(self, text: str) -> None:
        p = self.doc.add_paragraph()
        fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=12, after=6)
        r = p.add_run(text)
        set_run_font(r, size=13, bold=True, color=NAVY)

    def body(self, text: str) -> None:
        p = self.doc.add_paragraph()
        fmt(p, first=0.75)
        r = p.add_run(text)
        set_run_font(r, size=11, color=GRAY)

    def table(self, rows: list[list[str]], header=True) -> None:
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = t.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
                r = p.add_run(val)
                set_run_font(r, size=10, bold=(header and i == 0))
                if header and i == 0:
                    shade_cell(cell, "E8EEF7")
                elif i == len(rows) - 1:
                    shade_cell(cell, "EAF3EA")
        cap = self.doc.add_paragraph()
        fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
        r = cap.add_run("Tablo 1. Walk-forward karşılaştırma (sunum sırası) ve kilitli test.")
        set_run_font(r, size=9, italic=True, color=GRAY)

    def table2(self, rows: list[list[str]], caption: str) -> None:
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = t.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
                r = p.add_run(val)
                set_run_font(r, size=10, bold=(i == 0))
                if i == 0:
                    shade_cell(cell, "E8EEF7")
        cap = self.doc.add_paragraph()
        fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
        r = cap.add_run(caption)
        set_run_font(r, size=9, italic=True, color=GRAY)


def build() -> None:
    d = Doc()
    d.title("Elektrik Piyasası Fiyat Tahmini")
    d.body(
        "İspanya gün öncesi elektrik piyasasında saatlik fiyatı, sızıntısız bir zaman serisi "
        "hattıyla tahmin ettim. Karşılaştırma ölçütü, dünün aynı saatini kopyalayan naif modeldir. "
        "Sunumdaki model sırası sabittir: Ridge, LightGBM, XGBoost, ARIMA, Ridge+B+AR(1)."
    )
    d.body(
        "Elektrik fiyat tahmini, ilk bakışta talep ve üretim dengesine indirgenebilir. Yük artınca "
        "fiyat yükselir, rüzgâr üretimi güçlenince fiyat baskılanır; bir modelin bu ilişkiyi öğrenmesi "
        "beklenir. Bu çalışmada gördüğüm asıl mesele algoritma seçimi değildi. Asıl mesele, teslim "
        "saati geldiğinde elde gerçekten hangi bilginin bulunduğunu ayırmaktı."
    )
    d.body(
        "Çalışmayı İspanya gün öncesi elektrik piyasası üzerinde yürüttüm. Hedef değişken price day "
        "ahead, birimi €/MWh. Bu, teslim saatinden önce ihaleyle oluşan saatlik fiyattır. Veri "
        "setinde ayrıca price actual yer alıyor; gerçekleşmiş piyasa tarafına daha yakın bir seri. "
        "Bu kolonu özellik olarak kullanmadım. Kullanılsaydı doğrulama skoru yükselirdi, ancak "
        "skorun operasyonel karşılığı kalmazdı."
    )

    d.h("Problem tanımı")
    d.body(
        "Gün öncesi piyasada her teslim saati için ayrı bir fiyat oluşur. Tahmin, o saatin gerçekleşmiş "
        "üretim, yük ve hava bilgisi görülmeden üretilmek zorundadır. Bu bilgiler teslim anında henüz "
        "elde değildir; sızıntısız bir kurulumda da kullanılamaz."
    )
    d.body(
        "Araştırma sorusunu buna göre kurdum: teslim saati t gelmeden önce yayınlanmış ihale fiyatları, "
        "yük ve yenilenebilir gün öncesi tahminleri ile geçmiş üretim ve hava gözlemleri kullanılarak, "
        "dünün aynı saatini kopyalayan naif modelden daha düşük hata elde edilebilir mi? Naif model ilk "
        "bakışta zayıf bir temel gibi durur. Elektrik piyasasında durum farklıdır. Dünün 18.00 fiyatı, "
        "bugünün 18.00 fiyatına çoğu günde yeterince yakındır. Lag-24’ü geçemeyen bir sistemin pratik "
        "katkısı sınırlı kalır. Bu nedenle tüm resmi karşılaştırmayı bu tabana bağladım."
    )

    d.h("Veri seti")
    d.body(
        "Kaynak iki ham CSV dosyasından oluşuyor. Enerji dosyası üretim kırılımlarını, yükü, gün öncesi "
        "tahminleri ve fiyatları içeriyor; 35.064 saatlik kayıt. Hava dosyası Madrid, Barcelona, Bilbao, "
        "Sevilla ve Valencia için aynı saat diliminde gözlem taşıyor. Kapsam, Aralık 2014’ün son "
        "saatinden 2018 yıl sonuna kadardır. Seri saatlik, UTC zaman damgalı ve kopuksuzdur. Hedef 2,06 "
        "ile 101,99 €/MWh arasında değişiyor; boş fiyat kaydı yoktur."
    )
    d.body(
        "Üretim kolonlarında boşluk vardı. Bu boşlukları sıfırla doldurmadım. En fazla üç saatlik iç "
        "boşlukları zamana göre tamamladım, daha uzun ve kenar boşluklarını eksik bıraktım. Sıfır, veri "
        "yok anlamına gelmez; modele öyle yedirilirse hata gizlenmiş olur. Ham dosyaların üzerine "
        "yazmadım. Temizlik işlemlerini kopyalar üzerinde yürüttüm."
    )

    d.h("Sızıntı kontrolü ve zaman serisi bölmesi")
    d.body(
        "Zaman serisinde rastgele eğitim–test bölmesi, gelecek saati geçmişe karıştırır. Model bu "
        "sızıntıyı öğrenir; elde edilen skor yanıltıcı olur. Seriyi karıştırmadım. Bölmeyi kronolojik "
        "%70 / %15 / %15 olarak kestim. Eğitim dönemi 2014 sonundan 2017 Ekim’ine, doğrulama 2018 "
        "Mayıs’ına, test 2018’in kalanına uzanıyor."
    )
    d.body(
        "Test setini model seçiminde, hiperparametre ayarında, eşik belirlemede veya artık "
        "düzeltmesinde kullanmadım. Testi dondurulmuş hâliyle, hattın en sonunda bir kez skorladım. "
        "Hedef gecikmelerini t-24, t-48 ve t-168 ile sınırladım; t-1 yoktur. price actual hiç "
        "kullanılmadı. Yüksek fiyat oranlarını, önce 24 saat kaydırılmış seri üzerinde hesapladım."
    )

    d.h("Özellik mühendisliği")
    d.body(
        "Hattın sonunda 184 güvenli kolon oluştu. Fit anında üç kolon daha eklendi: son 7, 14 ve 30 "
        "günde fiyatın belirli bir eşiğin üzerinde kaldığı saatlerin oranı. Eşik, testten değil, "
        "geliştirme döneminin 75. yüzdeliğinden geldi. Bu üçlüye METHOD_B adını verdim. İşlevi "
        "sadedir: pahalı saatlerin geçmişte kümelenip kümelenmediğine bakan, nedensel bir frekans ölçüsü."
    )
    d.body(
        "Kolon grupları takvim, gün öncesi yük, güneş ve rüzgâr tahminleri, geçmiş fiyat, geçmiş üretim "
        "ve gecikmeli hava gözlemlerinden oluşuyor. Takvim, Madrid saatine çevrilmiş saat, haftanın "
        "günü, ay ve döngüsel kodları içeriyor."
    )

    d.h("Model seçimi ve doğrulama")
    d.body(
        "Eğitim ve doğrulamayı birleştirip walk-forward doğrulama uyguladım. Dört genişleyen pencere "
        "kullandım. Her katmanda eksik değer doldurma ve ölçekleme yalnızca o katmanın eğitim bloğunda "
        "fit edildi. Test dosyası bu aşamada açılmadı. Karşılaştırmayı sunum sırasıyla yürüttüm."
    )
    d.body(
        "Birinci model Ridge’dir. Ceza katsayısı walk-forward ızgarasında 0,001 seçildi. 184 güvenli "
        "özellikte ortalama MAE 5,80 oldu. Ridge çözümünü NumPy’de kapalı formda aldım."
    )
    d.body(
        "İkinci model LightGBM, üçüncü model XGBoost’tur. İkisi de aynı 184 özellikte, sızıntısız "
        "kurulumda denendi. Beklentim ağaçların öne geçmesi yönündeydi. Geçmediler. LightGBM 5,92, "
        "XGBoost 5,95. Ağaçlar daha esnektir; bu veri, sızıntı kesildikten sonra o esnekliği ödüllendirmedi."
    )
    d.body(
        "Dördüncü adım ARIMA’dır. Ham fiyata 184 özelliği yutan klasik tek değişkenli ARIMA kurmadım. "
        "Ridge ve METHOD_B çok değişkenli kısmı aldıktan sonra artığa ARIMA(1,0,0), yani AR(1) "
        "oturtuldu. Protokol gün öncesine uygundur: 24 saatlik blok tahmin, gün bitince artıklar "
        "güncellenir. Walk-forward MAE 4,53’e indi. Mevsimsel ARIMA daha çok parametreyle AIC’de "
        "önde göründü; tahmin hatasında AR(1)’den kötü kaldı. Bu yüzden teslimata sade AR(1) alındı."
    )
    d.body(
        "Beşinci ve kilitli model Ridge+B+AR(1)’dir. Yani Ridge (α = 0,001), METHOD_B’nin üç sıklığı, "
        "geliştirme artıklarından dondurulmuş +1,47 €/MWh ekleme ve artığa AR(1). Walk-forward MAE "
        "4,53’tür. Test bu bileşim seçildikten sonra bir kez skorlandı."
    )

    d.table(
        [
            ["Sıra", "Model", "Walk-forward MAE", "RMSE", "R²", "Test MAE"],
            ["1", "Ridge", "5,80", "7,41", "0,622", "—"],
            ["2", "LightGBM", "5,92", "7,85", "0,606", "—"],
            ["3", "XGBoost", "5,95", "7,91", "0,601", "—"],
            ["4", "ARIMA", "4,53", "6,07", "0,748", "—"],
            ["5", "Ridge+B+AR(1)", "4,53", "6,07", "0,748", "3,99"],
        ]
    )

    d.h("Kilitli test sonuçları")
    d.body(
        "Kilitli test seti 5.260 saati kapsıyor; 2018 Mayıs sonundan yıl sonuna. Aşağıdaki tablo, "
        "dondurulmuş holdout üzerindeki resmi karşılaştırmadır. Test, Ridge+B+AR(1) seçildikten sonra "
        "bir kez skorlanmıştır."
    )
    d.table2(
        [
            ["Metrik", "Ridge+B+AR(1)", "Naif Lag-24"],
            ["MAE", "3,99", "6,05"],
            ["RMSE", "5,88", "9,03"],
            ["R²", "0,67", "0,22"],
            ["sMAPE", "%7,31", "%11,68"],
            ["Sapma", "+1,42", "≈ 0"],
        ],
        "Tablo 2. Kilitli test performansı (5.260 saat, 2018 Mayıs–Aralık).",
    )
    d.body(
        "Mutlak hata yaklaşık yüzde 34 daha düşüktür. Bu metrikler tek başına yeterli değildir. Model "
        "sapması +1,42; yani ortalama olarak fiyat hâlâ yukarı yazılmıştır. Naif model neredeyse "
        "yansızdır. Düşük MAE, yansız bir model anlamına gelmez. Fiyat düzeyinde de kayma vardır: "
        "geliştirme ortalaması yaklaşık 47,9, test 61,1 €/MWh. Test dilimi hem daha pahalı hem daha "
        "az oynaktır."
    )

    d.h("Hata analizi")
    d.body(
        "Geliştirme katmanlarında Ridge+METHOD_B, pahalı saatleri sistematik olarak eksik tahmin "
        "ediyordu. P75 üstü sapma yaklaşık −5,9, P90 üstünde −8,2 idi. METHOD_B bu sapmayı yumuşattı, "
        "ortadan kaldırmadı. AR(1) katmanı ortalama hatayı indirdi; kuyruk sapmasını sihirli biçimde "
        "silmedi. Kilitli testte işaret döndü. Genel sapma pozitif; P75, P90 ve P95 dilimleri de "
        "pozitiftir. Doğrulama döneminde gözlediğim eksik tahmini çözdüğümü söyleyemem."
    )
    d.body(
        "Test skorunu sonradan iyileştirme baskısı oluşur. Eşik kaydırılabilir, model ailesi "
        "değiştirilebilir. MAE düşer; fakat elde kalan şey bağımsız bir değerlendirme olmaktan çıkar. "
        "Ridge+B+AR(1) walk-forward’da seçildi, test bir kez skorlandı. Panodaki rakam da bu kilitli "
        "değerlendirmedir; sistem yeniden eğitilmez."
    )

    d.h("Açıklanabilirlik")
    d.body(
        "Kilitli modelin omurgası doğrusal olduğu için TreeSHAP kullanmadım. Walk-forward katmanlarında "
        "Ridge için tam doğrusal SHAP hesapladım. En güçlü öngörü sinyali geçmiş fiyattır: t-24 ve "
        "t-48. Bunu toplam yük tahmini izler. Rüzgârın yük içindeki payı, güneş tahmini ve yenilenebilir "
        "tahmin toplamı da katkı verir. Bu, rüzgârın fiyatı düşürdüğüne dair nedensel bir kanıt "
        "değildir. AR(1) katmanı tek bir düzey düzeltmesidir; özellik bazlı SHAP terimi değildir."
    )
    d.body(
        "Takvim tarafında month ile day_of_year neredeyse aynı bilgiyi taşır. Katsayıları büyük ve zıt "
        "çıkabilir. Bunu ayrı sürücüler gibi okumak yanlıştır. SHAP, tahminin içini açar; nedensellik "
        "kanıtı sunmaz."
    )

    d.h("Operasyonel 24 saatlik tahmin")
    d.body(
        "Mevcut dosyalar ve kilitli 187 kolonlu Ridge+METHOD_B omurgası ile, D-1 öğlen civarı bir çıkış "
        "anında 24 saatlik operasyonel tahmin üretilemez. 106 kolon güvenli, 6 kolonun yayın saati "
        "dosyada doğrulanmamış, 75 kolon o anda henüz gözlenmemiştir. Katı üretim yolunda eksik kolonu "
        "sıfırla doldurmadım. Tahmin boş kalır. Karar kilitlidir: 24 saatlik üretim tahmini hazır değildir."
    )

    d.h("Karar panosu ve tekrarlanabilirlik")
    d.body(
        "Sonuca salt okunur bir Streamlit panosu ekledim. Uygulamanın adı projenin adıyla aynıdır: "
        "Elektrik Piyasası Fiyat Tahmini. Pano model eğitmez, SHAP hesaplamaz, 24 saatlik üretim basmaz. "
        "Model karşılaştırması sayfası Ridge, LightGBM, XGBoost, ARIMA ve Ridge+B+AR(1) sırasını korur. "
        "Ham CSV’den kilitli modele giden adımlar ayrı betiklerde ve defterde durur. Test seti en sonda, "
        "bir kez skorlanır."
    )

    d.h("Değerlendirme ve sınırlılıklar")
    d.body(
        "Elektrik piyasasında fiyat tahmini, daha derin bir model seçmekten önce bilgi sınırını doğru "
        "kurma problemidir. t-1 gecikmesi eklendiğinde çoğu model iyileşir. Operasyonda o gecikme yoksa "
        "iyileşme de yoktur. LightGBM ve XGBoost bu sette Ridge’i geçmedi. ARIMA katmanı, Ridge’in "
        "bıraktığı düzeyi dünden yarına taşıdığı için işe yaradı; ham seriyi tek başına açıklayan bir "
        "sihir değildi."
    )
    d.body(
        "Kilitli holdout’ta naif modeli geçtim; MAE 3,99’a indi. Buna karşılık sapma kaldı, işaret bile "
        "değişti. Yüksek fiyat rejimini çözdüğümü iddia etmiyorum. 24 saatlik üretim teslimi de yoktur. "
        "Yakıt, arıza ve sınır ötesi kısıt gibi dış değişkenler eksiktir. Tek bir 2018 dilimi, bir "
        "piyasanın tamamını temsil etmez."
    )
    d.body(
        "Elimde kalan çerçeve nettir. Kronolojik, sızıntısız, walk-forward ile seçilmiş ve testte bir "
        "kez değerlendirilmiş bir hat. Sunum sırası Ridge, LightGBM, XGBoost, ARIMA, Ridge+B+AR(1)’dir. "
        "Kilitli teslimat sonuncusudur. Gün öncesi fiyat tahmini için önce ilgili saatte elde bulunan "
        "bilgi yazılmalıdır; model ondan sonra kurulmalıdır."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    d.doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    build()
