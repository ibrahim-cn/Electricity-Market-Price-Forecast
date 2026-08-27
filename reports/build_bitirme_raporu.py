#!/usr/bin/env python3
"""softITo bitirme projesi raporu (örnek PDF şablonuna uygun Word belgesi)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "reports" / "bitirme_figures"
OUT = ROOT / "reports" / "Bitirme_Projesi_Raporu.docx"
OUT_ROOT = ROOT / "Bitirme_Projesi_Raporu.docx"

NAVY = RGBColor(0x1F, 0x2A, 0x44)
BLUE = RGBColor(0x1F, 0x4E, 0x9B)
GRAY = RGBColor(0x44, 0x44, 0x44)


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(p, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=8, first_line=0, line=1.15):
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    if first_line:
        pf.first_line_indent = Cm(first_line)


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, size=10, color=GRAY)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def shade_cell(cell, fill="F2F4F8"):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


class Report:
    def __init__(self):
        self.doc = Document()
        self._setup()
        self.fig_n = 0

    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)
        add_page_number(sec)
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    def p(self, text, *, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=8, first_line=0.75, color=None):
        para = self.doc.add_paragraph()
        set_paragraph_format(para, align=align, space_before=space_before, space_after=space_after, first_line=first_line)
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
        return para

    def mixed(self, parts, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, first_line=0.75, size=12):
        para = self.doc.add_paragraph()
        set_paragraph_format(para, align=align, space_after=space_after, first_line=first_line)
        for text, kwargs in parts:
            run = para.add_run(text)
            set_run_font(run, size=size, **kwargs)
        return para

    def h1(self, text):
        para = self.doc.add_paragraph()
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=14, first_line=0, line=1.0)
        para.paragraph_format.page_break_before = True
        run = para.add_run(text.upper())
        set_run_font(run, size=16, bold=True, color=NAVY)

    def h2(self, text):
        para = self.doc.add_paragraph()
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=14, space_after=8, first_line=0, line=1.0)
        run = para.add_run(text)
        set_run_font(run, size=13, bold=True, color=NAVY)

    def h3(self, text):
        para = self.doc.add_paragraph()
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=6, first_line=0, line=1.0)
        run = para.add_run(text)
        set_run_font(run, size=12, bold=True, color=BLUE)

    def h4(self, text):
        para = self.doc.add_paragraph()
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4, first_line=0, line=1.0)
        run = para.add_run(text)
        set_run_font(run, size=12, bold=True)

    def bullets(self, items, *, bold_lead=True):
        for item in items:
            para = self.doc.add_paragraph(style="List Bullet")
            set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=4, first_line=0, line=1.15)
            if bold_lead and ":" in item:
                lead, rest = item.split(":", 1)
                r1 = para.add_run(lead + ":")
                set_run_font(r1, size=12, bold=True)
                r2 = para.add_run(rest)
                set_run_font(r2, size=12)
            else:
                r = para.add_run(item)
                set_run_font(r, size=12)

    def numbered(self, items):
        for item in items:
            para = self.doc.add_paragraph(style="List Number")
            set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=4, first_line=0, line=1.15)
            if ":" in item:
                lead, rest = item.split(":", 1)
                r1 = para.add_run(lead + ":")
                set_run_font(r1, size=12, bold=True)
                r2 = para.add_run(rest)
                set_run_font(r2, size=12)
            else:
                r = para.add_run(item)
                set_run_font(r, size=12)

    def constraint(self, n, text):
        para = self.doc.add_paragraph()
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8, first_line=0)
        r1 = para.add_run(f"Kısıt {n}: ")
        set_run_font(r1, size=12, bold=True)
        r2 = para.add_run(text)
        set_run_font(r2, size=12)

    def fig(self, filename, caption, width=15.2):
        path = FIG / filename
        if not path.exists():
            self.p(f"[Şekil bulunamadı: {filename}]", italic=True, first_line=0, color=GRAY)
            return
        self.fig_n += 1
        para = self.doc.add_paragraph()
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2, first_line=0)
        para.add_run().add_picture(str(path), width=Cm(width))
        cap = self.doc.add_paragraph()
        set_paragraph_format(cap, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=12, first_line=0)
        run = cap.add_run(f"Şekil {caption}")
        set_run_font(run, size=11, italic=True)

    def table(self, headers, rows, col_widths=None):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            p = hdr[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h)
            set_run_font(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            shade_cell(hdr[i], "1F4E9B")
        for ri, row in enumerate(rows):
            cells = table.rows[ri + 1].cells
            fill = "F7F9FC" if ri % 2 == 0 else "FFFFFF"
            for ci, val in enumerate(row):
                cells[ci].text = ""
                p = cells[ci].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(str(val))
                set_run_font(r, size=10)
                shade_cell(cells[ci], fill)
        self.doc.add_paragraph()

    def code(self, text):
        para = self.doc.add_paragraph()
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=10, first_line=0, line=1.0)
        run = para.add_run(text)
        set_run_font(run, name="Consolas", size=10)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F4F6F8")
        shd.set(qn("w:val"), "clear")
        pPr.append(shd)

    def toc_line(self, left, right="", *, keep_with_next=True):
        para = self.doc.add_paragraph()
        set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=1, first_line=0, line=1.0)
        pf = para.paragraph_format
        pf.keep_together = True
        pf.keep_with_next = keep_with_next
        pf.widow_control = True
        pf.page_break_before = False
        pf.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
        is_chapter = left[:2] in ("1.", "2.", "3.", "4.", "5.") and not left.startswith(" ")
        r1 = para.add_run(left)
        set_run_font(r1, size=12, bold=is_chapter, color=NAVY if is_chapter else None)
        if right:
            r2 = para.add_run("\t" + right)
            set_run_font(r2, size=12)

    def cover(self):
        for _ in range(2):
            self.doc.add_paragraph()
        logo = FIG / "logo_softito.png"
        if logo.exists():
            p = self.doc.add_paragraph()
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, first_line=0)
            p.add_run().add_picture(str(logo), width=Cm(5.2))
        self.p("VERİ BİLİMİ BİTİRME PROJESİ RAPORU", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, space_before=10, space_after=10, color=NAVY)
        self.p("İspanya Gün Öncesi Elektrik Fiyat Tahmini", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, space_after=28)
        self.p("İbrahim Can", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, space_after=4)
        for _ in range(6):
            self.doc.add_paragraph()
        self.p("Ağustos, 2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, space_after=2)
        self.p("softITo Yazılım Bilişim Akademisi, İstanbul", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, space_after=0)
        new = self.doc.add_section()
        new.page_width = Cm(21.0)
        new.page_height = Cm(29.7)
        new.top_margin = Cm(2.5)
        new.bottom_margin = Cm(2.5)
        new.left_margin = Cm(2.5)
        new.right_margin = Cm(2.5)
        add_page_number(new)
        s0 = self.doc.sections[0]
        s0.footer.is_linked_to_previous = False
        s0.footer.paragraphs[0].clear()

    def toc(self):
        self.p("İÇİNDEKİLER", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=0, space_after=8, color=NAVY)
        items = [
            ("1. GİRİŞ", ""),
            ("     1.1. Projenin Amacı", ""),
            ("     1.2. Proje Kısıtları", ""),
            ("2. MATERYAL VE YÖNTEM", ""),
            ("     2.1. Veri Seti", ""),
            ("     2.2. Sızıntısız Zaman Serisi Modelleme", ""),
            ("     2.3. Ridge Regresyon Kavramı", ""),
            ("     2.4. Walk-Forward Doğrulama", ""),
            ("     2.5. Artık Düzeltmesi ve METHOD_B", ""),
            ("     2.6. ARIMA ve AR(1) Artık Düzeltmesi", ""),
            ("     2.7. SHAP ile Açıklanabilirlik", ""),
            ("     2.8. Streamlit Panosu", ""),
            ("3. UYGULAMA", ""),
            ("     3.1. Geliştirme Ortamı Kurulumu", ""),
            ("           3.1.1. Python Bağımlılıklarının Temini", ""),
            ("           3.1.2. Kurulum Adımları", ""),
            ("           3.1.3. Sürüm Kontrolü", ""),
            ("     3.2. Veri Hattı", ""),
            ("           3.2.1. Ham Verilerin Yüklenmesi", ""),
            ("           3.2.2. Temizleme ve Birleştirme", ""),
            ("           3.2.3. Özellik Mühendisliği", ""),
            ("           3.2.4. Kronolojik Bölme", ""),
            ("     3.3. Model Geliştirme ve Seçim", ""),
            ("           3.3.1. Temel (Baseline) Modeller", ""),
            ("           3.3.2. Walk-Forward Karşılaştırma", ""),
            ("           3.3.3. Ridge Alpha Seçimi", ""),
            ("           3.3.4. METHOD_B Seçimi", ""),
            ("           3.3.5. ARIMA / AR(1) Seçimi", ""),
            ("     3.4. Final Model ve Kilitli Test", ""),
            ("     3.5. Streamlit Panosunun Çalıştırılması", ""),
            ("           3.5.1. Pano Başlatma", ""),
            ("           3.5.2. Genel Bakış", ""),
            ("           3.5.3. Model Karşılaştırması", ""),
            ("           3.5.4. Model Performansı", ""),
            ("           3.5.5. Hata Analizi", ""),
            ("           3.5.6. Açıklanabilirlik", ""),
            ("           3.5.7. 24 Saatlik Tahmin ve Denetim", ""),
            ("     3.6. Test ve Doğrulama", ""),
            ("4. SONUÇ", ""),
            ("5. KAYNAKÇA", ""),
        ]
        for i, (left, right) in enumerate(items):
            self.toc_line(left, right, keep_with_next=(i < len(items) - 1))

    def body(self):
        self.h1("1. GİRİŞ")
        self.p(
            "Bu bölümde projenin amacı ve projede karşılaşılan kısıtlara yer verilmiştir."
        )

        self.h2("1.1. Projenin Amacı")
        self.p(
            "Bu projenin amacı, İspanya gün öncesi elektrik piyasasında saatlik fiyatı "
            "(price day ahead, €/MWh) sızıntısız ve kronolojik bir makine öğrenmesi hattı ile "
            "tahmin etmektir. Tahmin, teslim saati t gelmeden önce bilinebilecek bilgilerle "
            "üretilir: geçmiş ihale fiyatları, yük ve yenilenebilir gün-öncesi tahminleri, "
            "üretim ve hava durumu geçmişi ile takvim değişkenleri."
        )
        self.p(
            "Çalışmanın operasyonel sorusu şudur: Sızıntısız bir model, kilitli "
            "kronolojik bir holdout üzerinde Naive Lag-24 taban çizgisini geçebilir mi? "
            "Bu soru test seti açılmadan önce kapatılmış; test seti yalnızca bir kez "
            "skorlanmıştır. Sunum sırası sabittir: Ridge, LightGBM, XGBoost, ARIMA ve "
            "Ridge+B+AR(1). Teslimat, kilitli Ridge (α = 0,001) + METHOD_B + AR(1) modeli, "
            "doğrusal SHAP açıklanabilirliği ve salt-okunur bir Streamlit panosudur."
        )

        self.h2("1.2. Proje Kısıtları")
        self.constraint(
            1,
            "Rapor, İspanya gün öncesi elektrik fiyat tahminine odaklanır. Kripto varlık "
            "(ETH/USDT) tahmini, Reddit duygu analizi veya gerçek zamanlı price actual "
            "tahmini kapsam dışındadır.",
        )
        self.constraint(
            2,
            "Veri, 2014 sonu ile 2018 sonu arasındaki saatlik enerji ve beş şehirlik hava "
            "durumu kayıtlarıyla sınırlıdır. Canlı piyasa beslemesi, yakıt fiyatları, "
            "santral arızaları, politika şokları ve enterkonnektör kısıtları dosyada yoktur.",
        )
        self.constraint(
            3,
            "Test seti model seçimi, hiperparametre ayarı, eşik belirleme veya artık "
            "düzeltmesi için kullanılmamıştır. Model kilitlendikten sonra test bir kez "
            "değerlendirilmiş ve sonuçlar dondurulmuştur.",
        )
        self.constraint(
            4,
            "24 saatlik operasyonel üretim tahmini üretime hazır değildir. Gün-öncesi "
            "tahmin kolonlarının yayın saati dosyada bağımsız doğrulanmamıştır; katı "
            "(STRICT) yolda boş tahmin kasıtlıdır.",
        )
        self.constraint(
            5,
            "Bu çalışma öngörü başarımı ve ilişkilendirme raporlar. SHAP değerleri, "
            "katsayılar ve MAE kazancı nedensellik kanıtı olarak yorumlanmaz.",
        )

        self.h1("2. MATERYAL VE YÖNTEM")
        self.p(
            "Bu bölümde veri seti, sızıntı denetimi, özellik mühendisliği, doğrulama "
            "stratejisi, seçilen model ailesi ve açıklanabilirlik yöntemine yer verilmiştir."
        )

        self.h2("2.1. Veri Seti")
        self.p(
            "Proje iki ham CSV dosyası üzerine kuruludur. Ham dosyalar değiştirilmemiş; "
            "işleme kopyaları data/raw/ altında tutulmuştur."
        )
        self.table(
            ["Dosya", "Satır", "Kolon", "Rol"],
            [
                ["energy_dataset.csv", "35.064", "29", "Üretim, yük, gün-öncesi tahmin, fiyat"],
                ["weather_features.csv", "178.396", "17", "Madrid, Barcelona, Bilbao, Seville, Valencia hava"],
            ],
        )
        self.p("Veri seti şu bileşenleri içerir:", first_line=0)
        self.bullets(
            [
                "Hedef: Teslim saati t için yayınlanmış gün öncesi ihale fiyatı (price day ahead).",
                "Gün-öncesi tahminler: Toplam yük, kara rüzgârı ve güneş tahminleri.",
                "Gerçekleşmiş üretim: Fosil, nükleer, hidro, rüzgâr, güneş ve diğer kaynaklar.",
                "Gerçekleşmiş yük: total load actual (yalnızca gecikmeli kullanılır).",
                "Hava: Beş şehir için sıcaklık, nem, basınç, rüzgâr, yağış ve bulutluluk.",
                "Yasak kolon: price actual özellik matrisine hiç alınmaz.",
            ]
        )
        self.p("Çalışma mantığı:", first_line=0)
        self.numbered(
            [
                "Ham CSV’ler UTC saatlik ızgaraya hizalanır ve birleştirilir.",
                "Boş kolonlar düşülür; kısa iç boşluklar interpolasyonla doldurulur.",
                "Sızıntısız gecikmeler ve takvim kodları üretilir.",
                "Zaman sırası korunarak %70 / %15 / %15 bölünür.",
                "Geliştirme setinde walk-forward ile model seçilir; test bir kez skorlanır.",
            ]
        )
        self.table(
            ["Öğe", "Değer"],
            [
                ["Dönem", "2014-12-31 23:00 UTC → 2018-12-31 22:00 UTC"],
                ["Sıklık", "Saatlik, 0 eksik saat, monoton"],
                ["Hedef aralığı", "2,06 – 101,99 €/MWh"],
                ["Hedef eksik", "0"],
                ["Birleşik satır", "35.064"],
            ],
        )

        self.h2("2.2. Sızıntısız Zaman Serisi Modelleme")
        self.p(
            "Bilgi sınırı şudur: Teslim saati t için gün-öncesi fiyat, teslim anındaki "
            "gerçekleşmeler henüz yokken tahmin edilir. Rastgele karıştırma ve tesadüfi "
            "bölme kullanılmaz."
        )
        self.p("Zorunlu kontroller:", first_line=0)
        self.bullets(
            [
                "Kronolojik bölme: max(eğitim) < min(doğrulama) < min(test).",
                "Hedef t, t anının özelliği değildir.",
                "Aynı saatteki gerçekleşmiş üretim, yük ve hava kullanılmaz.",
                "Hedef gecikmeleri yalnızca t−24, t−48 ve t−168’dir; t−1 yoktur.",
                "184’lük ana matriste ham hedef üzerinde rolling pencere yoktur.",
                "Ön işleme (medyan impute + ölçekleme) her katta yalnızca o katın eğitim bloğuna oturtulur.",
                "Test; seçim, ayar, eşik veya artık eklemesi için açılmaz.",
            ]
        )

        self.h2("2.3. Ridge Regresyon Kavramı")
        self.p(
            "Ridge regresyon, en küçük kareler kestirimine L2 ceza ekleyen doğrusal bir "
            "yöntemdir. Çok sayıda ve birbirine bağlı özellik varken katsayıları şişmeden "
            "tutmaya yarar. Bu projede Ridge, kapalı form NumPy çözümü ile hesaplanır:"
        )
        self.code("(X'X + αI) w = X'y")
        self.p("Ridge şu özellikleri içerir:", first_line=0)
        self.bullets(
            [
                "Doğrusallık: Tahmin, ölçeklenmiş özelliklerin ağırlıklı toplamıdır.",
                "Düzenlileştirme: α büyüdükçe katsayılar küçülür; aşırı uyum riski azalır.",
                "Tekrarlanabilirlik: Karıştırma yoktur; aynı matris aynı ağırlığı verir.",
                "Açıklanabilirlik: Standartlaştırılmış katsayı ve tam doğrusal SHAP doğrudan yazılır.",
            ]
        )
        self.p("Çalışma mantığı:", first_line=0)
        self.numbered(
            [
                "Geliştirme katında medyan impute ve standart ölçekleme oturtulur.",
                "Seçilen α ile kapalı form Ridge çözülür.",
                "METHOD_B kesirleri yalnızca geçmiş y’den üretilir.",
                "Geliştirme artıklarından dondurulmuş expanding-historical ekleme uygulanır.",
                "Kilitli teslimatta bu tahmine bir de AR(1) artık düzeltmesi eklenir.",
            ]
        )
        self.p(
            "Kullanım alanları bu projede fiyat tahmini ile sınırlıdır. Ridge, sunumun "
            "birinci modelidir; LightGBM ve XGBoost aynı 184 özellikte walk-forward "
            "ortalama MAE’de Ridge’i geçmemiştir. Kilitlenen teslimat Ridge tek başına "
            "değil, METHOD_B ve AR(1) ile birleşimidir."
        )

        self.h2("2.4. Walk-Forward Doğrulama")
        self.p(
            "Tek bir 2017–2018 doğrulama kesiti tek rejimdir. Expanding pencereler, "
            "eğitim kökeni ileri kaydıkça hata kararlılığını ölçer. Dört kat, yalnızca "
            "eğitim + doğrulama (29.804 saat) üzerindedir; test.parquet yüklenmez."
        )
        self.table(
            ["Kat", "n_eğitim", "n_doğrulama", "Eğitim payı"],
            [
                ["1", "14.902", "2.980", "0,50"],
                ["2", "17.882", "2.980", "0,60"],
                ["3", "20.862", "2.981", "0,70"],
                ["4", "23.843", "5.961", "0,80"],
            ],
        )
        self.p(
            "Her katta max(eğitim zaman damgası) < min(doğrulama zaman damgası) şartı "
            "sağlanır. Bu yapı, model ailesi, α ve METHOD_B seçiminin doğru yeridir."
        )

        self.h2("2.5. Artık Düzeltmesi ve METHOD_B")
        self.p(
            "Geliştirme katlarında Ridge pahalı saatlerde sistematik eksik tahmin "
            "göstermiştir. Düzeltmeler yalnızca kat-eğitim artıklarıyla denenmiş; "
            "expanding-historical ekleme (τ = 720 saat) ortalama MAE’yi iyileştirmiştir."
        )
        self.p(
            "METHOD_B, geçmiş fiyattan t−24 kaydırılmış seri üzerinde 7 / 14 / 30 günlük "
            "yüksek fiyat kesirleri üretir. Eşik, kat eğitiminin (finalde geliştirme "
            "setinin) P75 değeridir (57,5825). Bu üç kolon model_features.parquet içinde "
            "saklanmaz; oturtma anında eklenir."
        )
        self.p(
            "Walk-forward’da METHOD_B ortalama MAE = 5,496022 ile seçilmiştir. Yüksek "
            "fiyat eksik tahmini geliştirmede tam çözülmemiştir (P75+ sapma −5,90). "
            "Bu geliştirme bulgusu test bulgusu gibi yeniden yazılmaz. METHOD_B, kilitli "
            "teslimatın orta katmanıdır; tek başına final model değildir."
        )

        self.h2("2.6. ARIMA ve AR(1) Artık Düzeltmesi")
        self.p(
            "Dördüncü sunum modeli ARIMA’dır. Ham fiyata tek değişkenli klasik ARIMA "
            "kurulmamıştır. Ridge ve METHOD_B çok değişkenli kısmı aldıktan sonra artığa "
            "ARIMA(1,0,0), yani AR(1) oturtulmuştur: hata_t ≈ φ × hata_{t−1}. φ, geliştirme "
            "artıklarının son 1.440 saati üzerinde yaklaşık 0,91 bulunmuştur."
        )
        self.p(
            "Protokol gün öncesine uygundur: 24 saatlik blok tahmin edilir, gün bitince "
            "o günün gerçek Ridge+METHOD_B artıkları eklenir. Gün içi lag-1 düzeltmesi "
            "gün-öncesi skor değildir. Walk-forward MAE 4,53’e inmiştir. Mevsimsel ARIMA "
            "AIC’de önde görünmüş, tahmin hatasında sade AR(1)’den kötü kalmıştır."
        )
        self.p(
            "Beşinci ve kilitli model Ridge+B+AR(1)’dir. Yani Ridge (α = 0,001), METHOD_B’nin "
            "üç sıklığı, dondurulmuş +1,47 €/MWh ekleme ve artığa AR(1). Walk-forward "
            "metrikleri dördüncü satırla aynıdır; çünkü kilitli teslimat o AR(1) katmanıdır. "
            "Test kolonu yalnızca bu beşinci satır için doldurulmuştur."
        )

        self.h2("2.7. SHAP ile Açıklanabilirlik")
        self.p(
            "Kilitli modelin omurgası doğrusaldır. TreeSHAP kullanılmaz. Tam doğrusal SHAP, "
            "kat-eğitim standardizasyonu sonrası φ_i = w_i · x_scaled,i olarak hesaplanır. "
            "METHOD_B eklemesi ve AR(1) düzeltmesi özellik bazlı SHAP terimi değildir."
        )
        self.p(
            "month ile day_of_year neredeyse eşdoğrusaldır; büyük ve zıt katsayıları "
            "bağımsız fiyat sürücüsü gibi okunmamalıdır. SHAP, kilitli tahminin "
            "açıklamasıdır; nedensellik iddiası değildir."
        )

        self.h2("2.8. Streamlit Panosu")
        self.p(
            "app.py salt-okunur bir panodur. Model eğitmez, α aramaz, METHOD_B değiştirmez, "
            "üretim tahmini basmaz. Kayıtlı CSV, parquet ve şekilleri gösterir. Menüde "
            "Genel bakış’tan sonra Model karşılaştırması gelir; sıra Ridge, LightGBM, "
            "XGBoost, ARIMA, Ridge+B+AR(1)’dir. Model durumu KİLİTLİ’dir; 24 saatlik "
            "tahmin üretime hazır değildir."
        )

        self.h1("3. UYGULAMA")
        self.p(
            "Bu bölümde ortam kurulumu, veri hattı, model seçimi, kilitli test, pano ve "
            "24 saatlik tahmin denetiminin uygulama adımlarına yer verilmiştir."
        )

        self.h2("3.1. Geliştirme Ortamı Kurulumu")
        self.h3("3.1.1. Python Bağımlılıklarının Temini")
        self.p(
            "Proje Python 3 ile çalışır. Gerekli paketler requirements.txt dosyasında "
            "tanımlıdır: numpy, pandas, pyarrow, scikit-learn, matplotlib ve streamlit."
        )
        self.h3("3.1.2. Kurulum Adımları")
        self.p("Proje kökünde aşağıdaki komut çalıştırılır:", first_line=0)
        self.code("python3 -m pip install -r requirements.txt")
        self.h3("3.1.3. Sürüm Kontrolü")
        self.p(
            "Kurulumun doğrulanması için Python sürümü ve streamlit komutu kontrol edilir. "
            "Pano, kilitli çıktıları okuyabildiği sürece model yeniden eğitilmez."
        )
        self.code("python3 --version\npython3 -m streamlit --version")

        self.h2("3.2. Veri Hattı")
        self.h3("3.2.1. Ham Verilerin Yüklenmesi")
        self.p(
            "energy_dataset.csv (35.064 satır, 29 kolon) ve weather_features.csv "
            "(178.396 satır, 17 kolon) okunur. Ham dosyalar üzerine yazılmaz."
        )
        self.h3("3.2.2. Temizleme ve Birleştirme")
        self.p(
            "Yüzde yüz boş iki kolon düşülür: generation hydro pumped storage aggregated "
            "ve forecast wind offshore eday ahead. En fazla 3 saatlik iç boşluklar zaman "
            "interpolasyonu ile doldurulur; daha uzun ve kenar boşluklar NaN bırakılır. "
            "Hedef interpolasyon yapılmaz. Hava tekrarları denetlenir; çıktı "
            "merged_energy_weather.parquet (35.064 × 103) dosyasıdır."
        )
        self.h3("3.2.3. Özellik Mühendisliği")
        self.p(
            "Sızıntısız 184 SAFE özellik üretilir. Hedef sonradan zaman damgası ile "
            "birleştirilir. METHOD_B’nin üç kesiri oturtma anında eklenir (187 kolon)."
        )
        self.table(
            ["Grup", "Sayı", "İçerik"],
            [
                ["Takvim", "20", "Europe/Madrid saat, gün, ay, çevrimsel kodlar"],
                ["Gün-öncesi tahmin", "6", "Yük / güneş / rüzgâr ve paylar"],
                ["Geçmiş hedef", "8", "t−24 / t−48 / t−168 ve bunların istatistikleri"],
                ["Geçmiş yük", "4", "Yük gecikmeleri ve lag-24 tahmin hatası"],
                ["Geçmiş üretim", "36", "Kaynak gecikmeleri ve toplamlar"],
                ["Geçmiş hava", "100", "Şehir hava gecikmeleri t−24 / t−168"],
                ["Hava agregatı", "10", "Ulusal ortalama / max yağış"],
            ],
        )
        self.h3("3.2.4. Kronolojik Bölme")
        self.p(
            "Karıştırma yoktur. 35.064 saatlik seri %70 / %15 / %15 kesilir. Hedef "
            "ortalaması eğitimde 46,85; doğrulamada 52,73; testte 61,14 €/MWh’dir. "
            "Bu rejim kayması, ilerideki sapma işareti değişiminin bağlamıdır."
        )
        self.table(
            ["Bölme", "Satır", "Başlangıç (UTC)", "Bitiş (UTC)"],
            [
                ["Eğitim", "24.544", "2014-12-31 23:00", "2017-10-19 14:00"],
                ["Doğrulama", "5.260", "2017-10-19 15:00", "2018-05-26 18:00"],
                ["Test", "5.260", "2018-05-26 19:00", "2018-12-31 22:00"],
            ],
        )
        self.fig("fig_split.png", "3.2.4.1. Kronolojik 70 / 15 / 15 bölme şeması")

        self.h2("3.3. Model Geliştirme ve Seçim")
        self.h3("3.3.1. Temel (Baseline) Modeller")
        self.p(
            "Tek doğrulama kesitinde Naive Lag-24 MAE = 8,30; Ridge (bu aşamadaki kaba "
            "ızgarada α = 0,1) MAE = 5,49 olmuştur. Test Naive Lag-24 değeri "
            "(MAE = 6,045924) önceden belirlenmiş taban çizgisidir; aile seçiminde "
            "kullanılmamıştır."
        )
        self.h3("3.3.2. Walk-Forward Karşılaştırma")
        self.p(
            "Karşılaştırma sunum sırasıyla raporlanır: Ridge, LightGBM, XGBoost, ARIMA, "
            "Ridge+B+AR(1). İlk üçü aynı 184 güvenli özelliktedir. ARIMA ve Ridge+B+AR(1) "
            "Ridge+METHOD_B artığına oturtulan AR(1)’dir; walk-forward MAE’leri aynıdır. "
            "Ağaçlar ortalama MAE’de Ridge’i geçmemiştir. HistGradientBoosting ve "
            "RandomForest da denenmiş, sunum sırasına alınmamıştır."
        )
        self.table(
            ["Sıra", "Model", "Ort. MAE", "Ort. RMSE", "Ort. R²", "Ort. sapma"],
            [
                ["1", "Ridge (α = 0,001)", "5,797", "7,414", "0,622", "−1,897"],
                ["2", "LightGBM", "5,917", "7,853", "0,606", "−3,153"],
                ["3", "XGBoost", "5,946", "7,906", "0,601", "−3,069"],
                ["4", "ARIMA (artık AR(1))", "4,530", "6,071", "0,748", "−0,713"],
                ["5", "Ridge+B+AR(1)", "4,530", "6,071", "0,748", "−0,713"],
            ],
        )
        self.fig("fig_walkforward.png", "3.3.2.1. Walk-forward model karşılaştırması (sunum sırası)")

        self.h3("3.3.3. Ridge Alpha Seçimi")
        self.p(
            "Aynı dört katta α ızgarası test görünmeden sabitlenmiştir: "
            "[0,001, 0,003, 0,01, 0,03, 0,1, 0,3, 1, 3, 10]. En düşük ortalama MAE "
            "α = 0,001’de (5,796612) elde edilmiş ve kilitlenmiştir."
        )
        self.h3("3.3.4. METHOD_B Seçimi")
        self.p(
            "Ridge α = 0,001 sabit tutulmuş; expanding-historical artık eklemesi "
            "üzerine yüksek fiyat yöntemleri yarıştırılmıştır. METHOD_B, 0,01 MAE "
            "eşikini aşarak seçilmiştir."
        )
        self.table(
            ["Yöntem", "Ort. MAE", "Ort. sapma", "P75+ sapma"],
            [
                ["CURRENT_BEST", "5,702", "−1,629", "−6,581"],
                ["METHOD_A", "5,530", "−1,180", "−6,030"],
                ["METHOD_B", "5,496", "−0,915", "−5,896"],
                ["METHOD_C", "5,522", "−1,215", "−5,908"],
            ],
        )
        self.fig("fig_method_b.png", "3.3.4.1. Yüksek fiyat stratejisi karşılaştırması")

        self.h3("3.3.5. ARIMA / AR(1) Seçimi")
        self.p(
            "Ridge α = 0,001 ve METHOD_B sabit tutulmuş; artığa AR(1), ARMA(1,1) ve "
            "mevsimsel ARIMA denenmiştir. Gün-öncesi 24 saatlik blok protokolünde en düşük "
            "walk-forward MAE AR(1)’de (4,530) elde edilmiştir. ARMA(1,1) 4,58; SAR(1) "
            "s = 24 4,71; AIC’nin tercih ettiği daha kalabalık SARIMA 4,87’dir. Teslimata "
            "tek parametreli AR(1) alınmıştır. Ham fiyata kurulan tek değişkenli ARIMA, "
            "sunumdaki dördüncü model değildir."
        )

        self.h2("3.4. Final Model ve Kilitli Test")
        self.p(
            "Kilitlenen yapı: Ridge (α = 0,001) + METHOD_B + AR(1). Geliştirme 29.804 "
            "satırdır. Dondurulmuş expanding-historical ekleme +1,471482; AR(1) φ ≈ 0,912’dir. "
            "Test parquet’i ağırlıklar, eşik, ekleme ve φ dondurulduktan sonra açılmıştır. "
            "5.260 satır skorlanmış, 0 satır düşülmüştür."
        )
        self.table(
            ["Metrik", "Ridge+B+AR(1)", "Naive Lag-24"],
            [
                ["MAE", "3,990091", "6,045924"],
                ["RMSE", "5,878929", "9,030375"],
                ["R²", "0,668961", "0,218923"],
                ["sMAPE", "%7,314412", "%11,675683"],
                ["Sapma", "+1,419419", "−0,003076"],
            ],
        )
        self.p(
            "MODEL_BEATS_NAIVE = EVET. MAE iyileşmesi yaklaşık %34’tür. Düşük MAE, "
            "modelin yansız olduğu anlamına gelmez. Naive bu holdout’ta sapmaya daha "
            "yakındır. Kilitli test sapması +1,42 (aşırı tahmin) iken geliştirmede sapma "
            "negatifti. P75+ / P90+ / P95+ sapma +0,84 / +0,92 / +0,56’dır. Yüksek fiyat "
            "eksik tahmini çözülmüş değildir; işaret dönmüştür. İki dönem tek hikâyede "
            "birleştirilmez."
        )
        self.fig("fig_test_mae.png", "3.4.1. Kilitli test MAE karşılaştırması")
        self.fig("fig_test_metrics.png", "3.4.2. Kilitli test MAE, RMSE ve sMAPE")
        self.fig("fig_actual_pred.png", "3.4.3. Günlük ortalama gerçek vs tahmin")
        self.fig("fig_scatter.png", "3.4.4. Gerçek vs tahmin saçılımı", width=11.0)
        self.fig("fig_residual.png", "3.4.5. Zaman içinde artık")

        self.h2("3.5. Streamlit Panosunun Çalıştırılması")
        self.h3("3.5.1. Pano Başlatma")
        self.p("Kilitli panoyu açmak için proje kökünde şu komut kullanılır:", first_line=0)
        self.code("python3 -m streamlit run app.py")
        self.p(
            "Pano http://localhost:8501 adresinde açılır. Yeniden eğitim yoktur; eksik "
            "çıktı dosyasında uyarı gösterilir."
        )
        self.h3("3.5.2. Genel Bakış")
        self.p(
            "Genel bakış sayfası kilitli KPI’ları, model kimliğini (Ridge+B+AR(1)) ve "
            "24 saatlik tahminin üretime hazır olmadığını gösterir."
        )
        self.fig("dash_overview.png", "3.5.2.1. Streamlit genel bakış ekranı")
        self.h3("3.5.3. Model Karşılaştırması")
        self.p(
            "Sunum sayfası beş modeli sabit sırada gösterir: Ridge, LightGBM, XGBoost, "
            "ARIMA, Ridge+B+AR(1). Sayılar walk-forward MAE’dir. Test MAE yalnızca beşinci "
            "satırda (3,99) doldurulur. ARIMA, ham fiyata tek değişkenli ARIMA değildir; "
            "Ridge+METHOD_B artığına ARIMA(1,0,0) oturtulmuştur."
        )
        self.fig("dash_compare.png", "3.5.3.1. Streamlit model karşılaştırması ekranı")
        self.h3("3.5.4. Model Performansı")
        self.p(
            "Performans sayfası kilitli test metriklerini, gerçek–tahmin zaman serisini "
            "ve artık dağılımını gösterir. Resmi kilitli skor yalnızca Ridge+B+AR(1) içindir."
        )
        self.fig("dash_perf.png", "3.5.4.1. Streamlit model performansı ekranı")
        self.fig("fig_mae_hour.png", "3.5.4.2. Saate göre MAE")
        self.fig("fig_mae_month.png", "3.5.4.3. Aya göre MAE")
        self.h3("3.5.5. Hata Analizi")
        self.p(
            "Geliştirmede pahalı saatlerde eksik tahmin vardı; kilitli testte sapma "
            "pozitife döndü. Bu işaret değişimi test görüldükten sonra düzeltilmedi, "
            "yalnızca raporlandı. P75+ sapma +0,84; P90+ +0,92; P95+ +0,56’dır. AR(1) "
            "ortalama hatayı indirmiş, kuyruk sapmasını silmemiştir."
        )
        self.fig("dash_error.png", "3.5.5.1. Streamlit hata analizi ekranı")
        self.fig("fig_quantile.png", "3.5.5.2. Fiyat dilimine göre hata")
        self.h3("3.5.6. Açıklanabilirlik")
        self.p(
            "Açıklanabilirlik sayfası kayıtlı doğrusal SHAP çıktılarını okur. İstikrarlı "
            "sinyaller arasında fiyat gecikmeleri, yük tahmini, rüzgârın yük payı ve "
            "üretim gecikmeleri vardır. Takvim çiftinin net katkısı yaklaşık sıfırdır. "
            "AR(1) katmanı tek bir düzey düzeltmesidir; özellik bazlı SHAP terimi değildir."
        )
        self.fig("dash_shap.png", "3.5.6.1. Streamlit açıklanabilirlik ekranı")
        self.fig("fig_shap_top.png", "3.5.6.2. En önemli özellikler (eşdoğrusal takvim çifti hariç)")
        self.fig("fig_shap_groups.png", "3.5.6.3. Özellik grubu |SHAP| toplamı")
        self.h3("3.5.7. 24 Saatlik Tahmin ve Denetim")
        self.p(
            "24 saatlik tahmin ancak gerekli her özellik tahmin başlangıcında biliniyorsa "
            "üretime hazır olur. D-1 ~12:00 CET kökeninde 187 kolondan 106 SAFE, 6 UNKNOWN "
            "ve 75 FORBIDDEN’tır. STRICT yol UNKNOWN/FORBIDDEN değerleri doldurmaz; "
            "boş y_pred kasıtlıdır. Varsayımsal senaryo üretime hazır değildir."
        )
        self.fig("dash_forecast.png", "3.5.7.1. 24 saatlik tahmin ekranı (üretime hazır değil)")
        self.fig("fig_availability.png", "3.5.7.2. Özellik uygunluğu denetimi")
        self.fig("dash_audit.png", "3.5.7.3. Veri bölmesi ve sızıntı denetimi ekranı")
        self.fig("dash_info.png", "3.5.7.4. Proje bilgisi ve aşama durumu")

        self.h2("3.6. Test ve Doğrulama")
        self.p("Doğrulama özeti:", first_line=0)
        self.bullets(
            [
                "Sızıntı denetimi: GEÇTİ.",
                "Test ile ayar / seçim: HAYIR.",
                "Karıştırma: Yok.",
                "Kilitli test satırı: 5.260; düşülen satır: 0.",
                "Naive’e göre MAE üstünlüğü: EVET (~%34).",
                "24 saatlik üretim tahmini: HAZIR DEĞİL.",
                "Tekrarlanabilirlik: Kapalı form Ridge + dondurulmuş AR(1) φ; random_state=42 ağaç karşılaştırmaları içindir.",
            ]
        )
        self.p(
            "Hattın sırası şöyledir: veri hazırlama → özellik mühendisliği → zaman "
            "bölmesi → baseline → walk-forward → Ridge ayarı → artık düzeltmesi → "
            "yüksek fiyat analizi → AR(1) artık düzeltmesi → SHAP → final model → kilitli test → 24s denetim → pano. "
            "Sonraki aşamalar kilitli skoru “iyileştirmek” için yeniden koşulmaz."
        )

        self.h1("4. SONUÇ")
        self.p(
            "Bu bitirme çalışmasında İspanya gün öncesi elektrik fiyatı, sızıntısız bir "
            "zaman serisi makine öğrenmesi hattı ile tahmin edilmiştir. Ham enerji ve "
            "hava verisinden kilitli holdout değerlendirmesine kadar tüm adımlar "
            "kronolojik tutulmuş; test seti seçim veya ayar için kullanılmamıştır."
        )
        self.p(
            "Seçilen model Ridge (α = 0,001) + METHOD_B + AR(1)’dir. Sunum sırası Ridge, "
            "LightGBM, XGBoost, ARIMA, Ridge+B+AR(1)’dir. Kilitli testte MAE 3,99; "
            "RMSE 5,88; R² 0,669 ve sMAPE %7,31 bulunmuş; Naive Lag-24 MAE’si 6,05 "
            "olarak kalmıştır. Model taban çizgisini yaklaşık %34 daha düşük MAE ile "
            "geçmiştir. Bununla birlikte düşük MAE, yansızlık veya yüksek fiyat hatasının "
            "çözüldüğü anlamına gelmez. Geliştirmede eksik tahmin, testte aşırı tahmine "
            "dönmüştür. AR(1) ortalama hatayı indirmiş; kuyruk sapmasını silmemiştir."
        )
        self.p(
            "Açıklanabilirlik, tam doğrusal SHAP ile yapılmıştır. Geçmiş fiyat "
            "gecikmeleri, yük ve yenilenebilir tahminleri ile üretim geçmişi istikrarlı "
            "öngörü sinyalleridir. Eşdoğrusal takvim kodları tek tek yorumlanmamıştır. "
            "Nedensellik iddiası yoktur."
        )
        self.p(
            "24 saatlik operasyonel tahmin, mevcut veri dosyasıyla üretime hazır "
            "değildir. Yayın saatleri doğrulanmadan ve yasak gecikmeler düşülmeden "
            "STRICT yol doldurulmuş tahmin basmaz. Bu karar kilitlidir."
        )
        self.p(
            "Sonuç olarak proje; sızıntı denetimli, walk-forward ile seçilmiş, bir kez "
            "kilitli holdout’ta skorlanmış ve salt-okunur panoda sunulmuş bir teslimattır. "
            "Kuramsal olarak zaman serisi sızıntı kontrolünü, pratik olarak da kilitli "
            "bir fiyat tahmin hattını belgelemektedir."
        )

        self.h1("5. KAYNAKÇA")
        refs = [
            "https://www.kaggle.com/datasets/nicholasjhana/energy-consumption-generation-prices-and-weather",
            "https://www.omie.es/en/market-results/daily-market",
            "https://www.statsmodels.org/stable/generated/statsmodels.tsa.arima.model.ARIMA.html",
            "https://scikit-learn.org/stable/modules/linear_model.html#ridge-regression",
            "https://shap.readthedocs.io/",
            "https://docs.streamlit.io/",
            "https://xgboost.readthedocs.io/",
            "https://lightgbm.readthedocs.io/",
            "Lundberg, S. M. ve Lee, S.-I. (2017). A Unified Approach to Interpreting Model Predictions. NeurIPS.",
            "Hyndman, R. J. ve Athanasopoulos, G. Forecasting: Principles and Practice.",
        ]
        for i, ref in enumerate(refs, 1):
            para = self.doc.add_paragraph()
            set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, first_line=0)
            r = para.add_run(f"[{i}]  {ref}")
            set_run_font(r, size=12, color=BLUE if ref.startswith("http") else None)
            if ref.startswith("http"):
                r.font.underline = True


def main():
    r = Report()
    r.cover()
    r.toc()
    r.body()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    r.doc.save(str(OUT))
    r.doc.save(str(OUT_ROOT))
    print("wrote", OUT)
    print("wrote", OUT_ROOT)


if __name__ == "__main__":
    main()
